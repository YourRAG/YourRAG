"""
File Parser Service for PDF and Word documents.
Extracts text content from PDF and DOCX files.
"""

import io
from typing import Optional
import ftfy


class FileParseError(Exception):
    """Exception raised when file parsing fails."""
    pass


class FileParser:
    """Service for parsing PDF and Word documents."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md', '.markdown'}
    
    @staticmethod
    def parse_pdf(file_bytes: bytes) -> str:
        """
        Extract text from PDF using PyMuPDF.
        
        Args:
            file_bytes: PDF file content as bytes
            
        Returns:
            Extracted text content
            
        Raises:
            FileParseError: If PDF parsing fails
        """
        try:
            import pymupdf
        except ImportError:
            try:
                import fitz as pymupdf
            except ImportError:
                raise FileParseError(
                    "PyMuPDF is not installed. Please install it with: pip install pymupdf"
                )
        
        try:
            doc = pymupdf.open(stream=file_bytes, filetype="pdf")
            text_parts = []
            
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text)
            
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            return FileParser._clean_text(full_text)
            
        except Exception as e:
            raise FileParseError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        """
        Extract text from DOCX using python-docx.
        
        Args:
            file_bytes: DOCX file content as bytes
            
        Returns:
            Extracted text content
            
        Raises:
            FileParseError: If DOCX parsing fails
        """
        try:
            from docx import Document
        except ImportError:
            raise FileParseError(
                "python-docx is not installed. Please install it with: pip install python-docx"
            )
        
        try:
            doc = Document(io.BytesIO(file_bytes))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_parts)
            return FileParser._clean_text(full_text)
            
        except Exception as e:
            raise FileParseError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def parse_text(file_bytes: bytes) -> str:
        """
        Parse plain text or markdown files.
        
        Args:
            file_bytes: Text file content as bytes
            
        Returns:
            Text content
        """
        try:
            # Try UTF-8 first
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Fallback to GBK for Chinese documents
                text = file_bytes.decode('gbk')
            except UnicodeDecodeError:
                # Last resort: ignore errors
                text = file_bytes.decode('utf-8', errors='ignore')
        
        return FileParser._clean_text(text)
    
    @staticmethod
    def parse_file(filename: str, file_bytes: bytes) -> str:
        """
        Parse file based on extension.
        
        Args:
            filename: Original filename with extension
            file_bytes: File content as bytes
            
        Returns:
            Extracted text content
            
        Raises:
            FileParseError: If file type is not supported or parsing fails
        """
        ext = FileParser._get_extension(filename)
        
        if ext == '.pdf':
            return FileParser.parse_pdf(file_bytes)
        elif ext == '.docx':
            return FileParser.parse_docx(file_bytes)
        elif ext == '.doc':
            raise FileParseError(
                "Legacy .doc format is not supported. Please convert to .docx format."
            )
        elif ext in {'.txt', '.md', '.markdown'}:
            return FileParser.parse_text(file_bytes)
        else:
            raise FileParseError(
                f"Unsupported file type: {ext}. "
                f"Supported types: {', '.join(sorted(FileParser.SUPPORTED_EXTENSIONS))}"
            )
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if file type is supported."""
        ext = FileParser._get_extension(filename)
        return ext in FileParser.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def _get_extension(filename: str) -> str:
        """Get lowercase file extension."""
        if '.' in filename:
            return '.' + filename.rsplit('.', 1)[-1].lower()
        return ''
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Clean extracted text.
        
        - Fix encoding issues using ftfy
        - Remove excessive whitespace
        - Normalize line breaks
        - Remove null characters
        """
        if not text:
            return ""
        
        # Fix encoding issues
        text = ftfy.fix_text(text)

        # Remove null characters
        text = text.replace('\x00', '')
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive blank lines (more than 2 consecutive)
        import re
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove excessive spaces
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Clean up lines
        lines = []
        for line in text.split('\n'):
            line = line.strip()
            lines.append(line)
        
        return '\n'.join(lines).strip()


# Convenience function
def parse_file(filename: str, file_bytes: bytes) -> str:
    """
    Parse file and extract text content.
    
    Args:
        filename: Original filename with extension
        file_bytes: File content as bytes
        
    Returns:
        Extracted text content
    """
    return FileParser.parse_file(filename, file_bytes)